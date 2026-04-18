"""
template_editor.py — Surgical .docx template editing.

When the user provides a .docx template, this module replaces only the
content of selected sections directly in the original file instead of
regenerating the whole document from scratch.  The cover page, TOC,
original styles and unchanged sections are fully preserved.

Public API
----------
find_section_positions(docx_path)           → list[dict]
build_combined_edit_prompt(...)             → str
parse_section_responses(llm_output)         → dict[str, str]
apply_section_edits(template, output, ...)  → Path
"""

from __future__ import annotations

import copy
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# ── Structured-output delimiters ──────────────────────────────────────

_SEP_START = "<<<SECTION:"
_SEP_END   = "<<<END_SECTION>>>"

# Word XML namespace
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ── Helpers ───────────────────────────────────────────────────────────

def _para_heading_level(para_el) -> int | None:
    """
    Return the heading level (1-9) of an XML paragraph element, or None.
    Matches both built-in (Heading1) and display-name styles (Heading 1).
    """
    pStyle = para_el.find(f".//{{{_W}}}pStyle")
    if pStyle is None:
        return None
    val = pStyle.get(f"{{{_W}}}val", "").lower().replace(" ", "")
    m = re.match(r"heading(\d+)$", val)
    return int(m.group(1)) if m else None


def _el_tag(el) -> str:
    tag = el.tag
    return tag.split("}")[-1] if "}" in tag else tag


# ── 1. Find section positions ─────────────────────────────────────────

def find_section_positions(docx_path: str) -> list[dict]:
    """
    Return an ordered list of heading positions in the .docx file.

    Each entry: ``{"title": str, "para_idx": int, "level": int}``
    where *para_idx* is the index within ``doc.paragraphs``.
    """
    doc = Document(str(docx_path))
    positions: list[dict] = []
    for idx, para in enumerate(doc.paragraphs):
        style = (para.style.name or "").lower()
        text  = para.text.strip()
        if not (style.startswith("heading") and text):
            continue
        # Extract level from style name
        m = re.search(r"(\d+)$", style)
        level = int(m.group(1)) if m else 1
        positions.append({"title": text, "para_idx": idx, "level": level})
    return positions


# ── 2. Build combined LLM prompt ──────────────────────────────────────

_FMT_ES = """\
Usa estas convenciones Markdown dentro de cada sección:
- `##` para el título de la sección (primera línea obligatoria)
- `###` para subsecciones cuando sea necesario
- Listas con `-` o `*` donde corresponda
- Tablas en sintaxis Markdown nativa con `|`
- Bloques de código con triple backtick y el lenguaje especificado
- Para diagramas: [DIAGRAM]\\n<código mermaid>\\n[/DIAGRAM]
"""

_FMT_EN = """\
Use these Markdown conventions inside each section:
- `##` for the section heading (mandatory first line)
- `###` for subsections when needed
- Lists with `-` or `*` where appropriate
- Native Markdown tables with `|`
- Code blocks with triple backtick and the language specified
- For diagrams: [DIAGRAM]\\n<mermaid code>\\n[/DIAGRAM]
"""


def build_combined_edit_prompt(
    repo_scan,
    sections_to_edit: list[str],
    all_sections: list[str],
    section_enrichments: dict | None,
    generator,
    output_lang: str = "es",
) -> str:
    """
    Build a single combined prompt for all sections to edit.

    The LLM is asked to return each section wrapped in
    ``<<<SECTION: name>>> … <<<END_SECTION>>>`` delimiters.
    """
    from . import ai_service

    lang = output_lang if output_lang in ("es", "en") else "es"

    repo_context = ai_service.build_repo_context(repo_scan)
    persona = (
        getattr(generator, f"PERSONA_{lang.upper()}", None)
        or getattr(generator, "PERSONA", "")
    )
    extra = (
        getattr(generator, f"EXTRA_INSTRUCTIONS_{lang.upper()}", None)
        or getattr(generator, "EXTRA_INSTRUCTIONS", "")
    )

    all_secs_text  = "\n".join(f"  - {s}" for s in all_sections)
    fmt_block      = _FMT_EN if lang == "en" else _FMT_ES

    # Per-section enrichment instructions
    _elabels = (
        {"table": "a Markdown table", "diagram": "a Mermaid diagram"}
        if lang == "en"
        else {"table": "una tabla Markdown", "diagram": "un diagrama Mermaid"}
    )
    _join = " and " if lang == "en" else " y "

    section_lines: list[str] = []
    for sec in sections_to_edit:
        enrichments = (section_enrichments or {}).get(sec, [])
        enrich_note = ""
        if enrichments:
            what = _join.join(_elabels[t] for t in enrichments if t in _elabels)
            if what:
                enrich_note = (
                    f" Include {what}." if lang == "en" else f" Incluye {what}."
                )
        if lang == "en":
            section_lines.append(f"  - **{sec}**: Generate full content.{enrich_note}")
        else:
            section_lines.append(f"  - **{sec}**: Redacta el contenido completo.{enrich_note}")

    sections_block = "\n".join(section_lines)

    if lang == "en":
        intro = (
            "You are a software documentation expert editing specific sections "
            "of an existing Word document template. Generate new content for each "
            "listed section based on the repository code below.\n\n"
        )
        persona_block = f"## Your role\n\n{persona}\n\n" if persona else ""
        extra_block   = f"## Specific instructions\n\n{extra}\n\n" if extra else ""
        ctx_block = (
            "## Document context\n\n"
            f"The full document contains these sections (for scope/context only):\n"
            f"{all_secs_text}\n\n"
        )
        task_block = (
            "## Sections to generate\n\n"
            f"{sections_block}\n\n"
        )
        fmt_hdr  = "## Formatting\n\n"
        out_block = (
            "## Output format — IMPORTANT\n\n"
            "For EACH section, wrap the output with these exact delimiters:\n\n"
            f"  {_SEP_START} <exact section name>>>\n"
            f"  ## <section heading>\n"
            f"  <section content>\n"
            f"  {_SEP_END}\n\n"
            "Rules:\n"
            f"- The section name between `{_SEP_START}` and `>>>` must match "
            "exactly as listed above.\n"
            "- First line of content must be `## <section name>`.\n"
            "- Output NOTHING outside the delimiters.\n"
            "- Base every statement on the actual repository code.\n"
            "- Write entirely in English.\n\n"
        )
    else:
        intro = (
            "Eres un experto en documentación de software. Estás editando secciones "
            "específicas de una plantilla de documento Word existente. Genera contenido "
            "nuevo para cada sección indicada basándote en el código del repositorio.\n\n"
        )
        persona_block = f"## Tu rol\n\n{persona}\n\n" if persona else ""
        extra_block   = f"## Instrucciones específicas\n\n{extra}\n\n" if extra else ""
        ctx_block = (
            "## Contexto del documento\n\n"
            f"El documento completo contiene estas secciones (solo para contexto):\n"
            f"{all_secs_text}\n\n"
        )
        task_block = (
            "## Secciones a generar\n\n"
            f"{sections_block}\n\n"
        )
        fmt_hdr  = "## Formato Markdown\n\n"
        out_block = (
            "## Formato de salida — IMPORTANTE\n\n"
            "Para CADA sección, envuelve la respuesta con estos delimitadores exactos:\n\n"
            f"  {_SEP_START} <nombre exacto de la sección>>>\n"
            f"  ## <título de la sección>\n"
            f"  <contenido de la sección>\n"
            f"  {_SEP_END}\n\n"
            "Reglas:\n"
            f"- El nombre entre `{_SEP_START}` y `>>>` debe coincidir exactamente "
            "con los nombres listados arriba.\n"
            "- La primera línea del contenido debe ser `## <nombre de sección>`.\n"
            "- NO escribas nada fuera de los delimitadores.\n"
            "- Basa cada afirmación en el código real del repositorio.\n"
            "- Redacta íntegramente en español.\n\n"
        )

    return (
        f"{intro}"
        f"{persona_block}"
        f"{fmt_hdr}{fmt_block}\n"
        f"{ctx_block}"
        f"{task_block}"
        f"{out_block}"
        f"{extra_block}"
        f"---\n\n{repo_context}"
    )


# ── 3. Parse LLM response ─────────────────────────────────────────────

def parse_section_responses(llm_output: str) -> dict[str, str]:
    """
    Extract ``{section_name: markdown_content}`` from the structured LLM output.

    Matches blocks delimited by ``<<<SECTION: name>>> … <<<END_SECTION>>>``.
    """
    result: dict[str, str] = {}
    pattern = re.compile(
        r"<<<SECTION:\s*(.+?)>>>\s*\n(.*?)<<<END_SECTION>>>",
        re.DOTALL,
    )
    for m in pattern.finditer(llm_output):
        name    = m.group(1).strip()
        content = m.group(2).strip()
        result[name] = content
    return result


# ── 4. Markdown → docx elements ───────────────────────────────────────

_INLINE_RE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|`([^`]+)`|([^*`]+)", re.DOTALL)


def _add_inline(para, text: str) -> None:
    for m in _INLINE_RE.finditer(text):
        b, it, code, plain = m.groups()
        if b is not None:
            r = para.add_run(b); r.bold = True
        elif it is not None:
            r = para.add_run(it); r.italic = True
        elif code is not None:
            r = para.add_run(code)
            r.font.name = "Courier New"; r.font.size = Pt(9)
        elif plain is not None:
            para.add_run(plain)


def _shade_para(para, hex_color: str = "F0F0F5") -> None:
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def _build_elements_from_markdown(scratch: Document, markdown: str) -> list:
    """
    Convert *markdown* to python-docx XML elements using a scratch document.

    Returns a list of lxml elements.  The first ``##`` heading line is
    stripped because the original template heading is kept in place.
    """
    lines = markdown.splitlines()
    n     = len(lines)
    elements: list = []
    i = 0
    first_heading_skipped = False

    while i < n:
        line     = lines[i]
        stripped = line.strip()

        # Skip [CAPTION:] tags
        if stripped.startswith("[CAPTION:") and stripped.endswith("]"):
            i += 1
            continue

        # Headings
        heading_m = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading_m:
            level = len(heading_m.group(1))
            text  = heading_m.group(2).strip()
            # Skip the first heading (== section title already in template)
            if not first_heading_skipped:
                first_heading_skipped = True
                i += 1
                continue
            style_map = {1: "Heading 2", 2: "Heading 2", 3: "Heading 3",
                         4: "Heading 4", 5: "Heading 4", 6: "Heading 4"}
            para = scratch.add_paragraph(style=style_map.get(level, "Heading 3"))
            para.clear()
            para.add_run(text)
            elements.append(para._element)
            i += 1
            continue

        # Fenced code block
        if line.startswith("```"):
            i += 1
            code_lines: list[str] = []
            while i < n and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            for cl in code_lines:
                para = scratch.add_paragraph(style="Normal")
                para.paragraph_format.left_indent  = Inches(0.3)
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after  = Pt(0)
                _shade_para(para)
                run = para.add_run(cl)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                elements.append(para._element)
            continue

        # Markdown table
        if "|" in line and stripped.startswith("|"):
            table_rows: list[str] = []
            while i < n and "|" in lines[i] and lines[i].strip().startswith("|"):
                table_rows.append(lines[i])
                i += 1
            data_rows = [r for r in table_rows if not re.match(r"^[\s|:\-]+$", r)]
            if data_rows:
                def _split_row(r: str) -> list[str]:
                    return [c.strip() for c in r.strip().strip("|").split("|")]
                cols = len(_split_row(data_rows[0]))
                tbl  = scratch.add_table(rows=len(data_rows), cols=cols)
                tbl.style = "Table Grid"
                for r_idx, raw in enumerate(data_rows):
                    cells = _split_row(raw)
                    for c_idx, cell_text in enumerate(cells[:cols]):
                        cell = tbl.cell(r_idx, c_idx)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        if r_idx == 0:
                            run = p.add_run(cell_text); run.bold = True
                        else:
                            _add_inline(p, cell_text)
                elements.append(tbl._tbl)
                spacer = scratch.add_paragraph()
                elements.append(spacer._element)
            continue

        # [DIAGRAM] block
        if stripped == "[DIAGRAM]":
            i += 1
            diag_lines: list[str] = []
            while i < n and lines[i].strip() != "[/DIAGRAM]":
                diag_lines.append(lines[i])
                i += 1
            i += 1  # skip [/DIAGRAM]
            lbl = scratch.add_paragraph()
            r = lbl.add_run("[Diagrama Mermaid:]"); r.italic = True; r.font.size = Pt(9)
            elements.append(lbl._element)
            for dl in diag_lines:
                para = scratch.add_paragraph(style="Normal")
                run = para.add_run(dl)
                run.font.name = "Courier New"; run.font.size = Pt(9)
                elements.append(para._element)
            continue

        # Unordered list
        if re.match(r"^\s*[-*]\s+", line):
            para = scratch.add_paragraph(style="List Bullet")
            _add_inline(para, re.sub(r"^\s*[-*]\s+", "", line))
            elements.append(para._element)
            i += 1
            continue

        # Ordered list
        if re.match(r"^\s*\d+\.\s+", line):
            para = scratch.add_paragraph(style="List Number")
            _add_inline(para, re.sub(r"^\s*\d+\.\s+", "", line))
            elements.append(para._element)
            i += 1
            continue

        # Blank line → spacer
        if not stripped:
            i += 1
            continue

        # Normal paragraph
        para = scratch.add_paragraph(style="Normal")
        _add_inline(para, line)
        elements.append(para._element)
        i += 1

    return elements


# ── 5. Apply surgical edits to the .docx ─────────────────────────────

def apply_section_edits(
    template_path: str,
    output_path: str,
    section_responses: dict[str, str],
    section_positions: list[dict],
) -> Path:
    """
    Replace section content in the original .docx template surgically.

    For each section in *section_responses*:
      - The heading paragraph is preserved.
      - Everything after the heading up to the next same-or-higher heading
        is deleted.
      - The LLM-generated markdown is converted to python-docx elements
        and inserted after the heading.

    Sections are processed in **reverse** para_idx order so that
    forward indices remain valid throughout the loop.
    """
    import shutil

    src  = Path(template_path)
    dest = Path(output_path)
    shutil.copy2(str(src), str(dest))

    doc  = Document(str(dest))
    body = doc.element.body

    # Index positions that have LLM responses
    edit_map: dict[int, dict] = {
        pos["para_idx"]: pos
        for pos in section_positions
        if pos["title"] in section_responses
    }
    if not edit_map:
        doc.save(str(dest))
        return dest

    # Process in REVERSE para_idx order to preserve forward indices
    for para_idx, pos_info in sorted(edit_map.items(), key=lambda x: x[0], reverse=True):
        title   = pos_info["title"]
        level   = pos_info["level"]
        content = section_responses.get(title, "").strip()
        if not content:
            continue

        # Re-fetch paragraphs (doc was mutated by previous iterations)
        current_paras = doc.paragraphs
        if para_idx >= len(current_paras):
            continue

        heading_el = current_paras[para_idx]._element

        # Map body children to find insertion/deletion range
        body_children  = list(body)
        try:
            h_body_idx = body_children.index(heading_el)
        except ValueError:
            continue

        # Find the end: next heading at same or higher (lower number) level
        end_body_idx = len(body_children)
        for j in range(h_body_idx + 1, len(body_children)):
            child = body_children[j]
            if _el_tag(child) == "p":
                child_level = _para_heading_level(child)
                if child_level is not None and child_level <= level:
                    end_body_idx = j
                    break

        # Identify the body-level sectPr so we never accidentally delete it.
        # It holds page margins and must always remain the last child of <w:body>.
        body_sect_pr = body.find(qn("w:sectPr"))

        # Delete existing content between heading and next heading,
        # but never touch the body sectPr.
        for child in body_children[h_body_idx + 1 : end_body_idx]:
            if child is body_sect_pr:
                continue
            body.remove(child)

        # Build replacement elements from markdown in a scratch document
        scratch   = Document()
        new_els   = _build_elements_from_markdown(scratch, content)

        # Insert after heading (before sectPr if it's still there)
        insert_at = h_body_idx + 1
        for offset, el in enumerate(new_els):
            body.insert(insert_at + offset, copy.deepcopy(el))

    # Re-anchor the body sectPr as the very last child of <w:body>.
    # OOXML requires this position; inserting paragraphs above may have
    # displaced it, which causes Word to ignore the stored margin settings.
    body_sect_pr = body.find(qn("w:sectPr"))
    if body_sect_pr is not None:
        body.remove(body_sect_pr)
        body.append(body_sect_pr)

    doc.save(str(dest))
    return dest
