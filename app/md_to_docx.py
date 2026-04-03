"""
md_to_docx.py — Converts a Markdown string to a Word (.docx) document.

Public API
----------
convert(markdown, output_path, *, doc_type, primary_color, secondary_color) -> Path
    Parse *markdown* and write a styled .docx to *output_path*.
    Returns the resolved Path of the written file.

Colour roles
------------
  primary_color   → H1, H2, table header background, footer title, header text
  secondary_color → H3+, inline/block code font

Document structure
------------------
  Cover page    → title centred (vert + horiz), subtitle, month/year, rights
  TOC page      → "Índice" (technical/user_manual) or "Agenda" (executive)
                  Clickable entries jump to heading bookmarks in the body
  Header        → placeholder logo (left) + "Autor(a):" / "Draft generado…" (right)
  Footer        → project title · "Todos los derechos reservados, Mes Año" (left)
  Margins       → 1.27 cm on all sides (page + header/footer distance)

Raises
------
RuntimeError — if the output directory does not exist or cannot be written.
"""

from __future__ import annotations

import copy
import io
import itertools
import re
import struct
import zlib
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm


# ── Colour defaults ───────────────────────────────────────────────────

_DEFAULT_PRIMARY   = "#2D3A8C"
_DEFAULT_SECONDARY = "#287A5F"

_COLOR_GRAY  = RGBColor(0x60, 0x60, 0x60)
_COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_COLOR_CODE_BG = "F0F0F5"

# Counter for unique SVG part names within the OPC package
_svg_counter = itertools.count(1)

# Namespace constants used for SVG embedding
_NS_ASVG = "http://schemas.microsoft.com/office/drawing/2016/SVG/main"
_NS_R    = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_NS_A    = "http://schemas.openxmlformats.org/drawingml/2006/main"
_REL_IMAGE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"


def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _hex_to_str(hex_color: str) -> str:
    """Return uppercase hex without '#', e.g. '2D3A8C'."""
    return hex_color.lstrip("#").upper()


# ── Localisation ──────────────────────────────────────────────────────

_MONTHS_ES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


def _month_year_es() -> str:
    now = datetime.now()
    return f"{_MONTHS_ES[now.month - 1].capitalize()} {now.year}"


# ── Placeholder logo PNG ──────────────────────────────────────────────

def _make_placeholder_png(width: int = 120, height: int = 50) -> bytes:
    """Build a minimal solid-colour PNG in memory (no Pillow needed)."""
    def chunk(tag: bytes, data: bytes) -> bytes:
        crc = zlib.crc32(tag + data) & 0xFFFFFFFF
        return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", crc)

    signature = b"\x89PNG\r\n\x1a\n"
    ihdr = chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
    row  = bytes([210, 215, 230] * width)
    raw  = b"".join(b"\x00" + row for _ in range(height))
    idat = chunk(b"IDAT", zlib.compress(raw, 6))
    iend = chunk(b"IEND", b"")
    return signature + ihdr + idat + iend


# ── Low-level XML helpers ─────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_para_shading(para, hex_color: str) -> None:
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def _get_or_add_tblPr(tbl_el) -> object:
    """Return the <w:tblPr> child of a CT_Tbl element, creating it if absent."""
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)
    return tblPr


def _remove_table_borders(table) -> None:
    tblPr = _get_or_add_tblPr(table._tbl)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def _set_cell_width(cell, width_twips: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(width_twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_cell_valign(cell, val: str = "center") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(old)
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), val)
    tcPr.append(vAlign)


def _center_table(tbl) -> None:
    tblPr = _get_or_add_tblPr(tbl._tbl)
    for old in tblPr.findall(qn("w:jc")):
        tblPr.remove(old)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tblPr.append(jc)


def _is_light_color(hex_color: str) -> bool:
    """Return True if the colour is light enough that dark text is more readable."""
    h = hex_color.lstrip("#")
    r = int(h[0:2], 16)
    g = int(h[2:4], 16)
    b = int(h[4:6], 16)
    return (0.299 * r + 0.587 * g + 0.114 * b) / 255.0 > 0.5


def _add_page_break(doc: Document) -> None:
    """Insert an explicit page break paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run  = para.add_run()
    br   = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# ── Bookmark helpers ──────────────────────────────────────────────────

def _add_bookmark(para, anchor_id: str, bk_id: int) -> None:
    """
    Wrap the content of *para* in a named bookmark so that internal
    hyperlinks (w:hyperlink w:anchor) can jump to it.
    """
    p = para._p

    bkStart = OxmlElement("w:bookmarkStart")
    bkStart.set(qn("w:id"),   str(bk_id))
    bkStart.set(qn("w:name"), anchor_id)

    # Insert right after <w:pPr> (if any), before the first run
    pPr = p.find(qn("w:pPr"))
    pos = (list(p).index(pPr) + 1) if pPr is not None else 0
    p.insert(pos, bkStart)

    bkEnd = OxmlElement("w:bookmarkEnd")
    bkEnd.set(qn("w:id"), str(bk_id))
    p.append(bkEnd)


def _add_toc_hyperlink(para, text: str, anchor: str,
                        color_hex: str, half_pts: int, bold: bool = False) -> None:
    """
    Append an internal hyperlink run to *para*.
    *half_pts* is the font size in half-points (e.g. 22 = 11 pt).
    """
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("w:anchor"), anchor)

    r    = OxmlElement("w:r")
    rPr  = OxmlElement("w:rPr")

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color_hex)
    rPr.append(color_el)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(half_pts))
    rPr.append(sz)

    if bold:
        rPr.append(OxmlElement("w:b"))

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    r.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)

    hl.append(r)
    para._p.append(hl)


def _add_caption(doc: Document, elem_type: str, number: int,
                 description: str, anchor_id: str, bk_id: int) -> None:
    """
    Add a styled caption paragraph below a table, code block or diagram.

    Format:  <elem_type> <number>,  <description>
    Example: Tabla 3,  Dependencias principales del proyecto
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(8)
    run = para.add_run(f"{elem_type} {number},  {description}")
    run.font.size      = Pt(9)
    run.font.italic    = True
    run.font.color.rgb = _COLOR_GRAY
    _add_bookmark(para, anchor_id, bk_id)


# ── Element pre-scan ──────────────────────────────────────────────────

def _prescan_elements(
    lines: list[str],
    title_line_idx: int,
    start_bk_id: int,
) -> tuple[dict, list, int]:
    """
    Scan *lines* in document order to assign per-type sequential caption numbers
    to every table, code block and diagram.

    Each element type has its own independent counter (Tabla 1…N, Código 1…N,
    Diagrama 1…N) so the numbers never overlap across types.

    Returns
    -------
    caption_map : dict[int, tuple]
        Keyed by the *element start line index*.
        Value: (elem_type, description, number, anchor_id, bk_id)
    table_toc_entries : list of (number, description, anchor_id)
        Only Table entries, used to build the Índice de Tablas in the TOC.
    next_bk_id : int
        Next available bookmark ID (continues from start_bk_id).
    """
    caption_map: dict[int, tuple] = {}
    table_toc_entries: list[tuple] = []

    table_num   = 0
    code_num    = 0
    diagram_num = 0
    bk_id       = start_bk_id
    pending_desc: str | None = None   # description from the nearest [CAPTION:]

    i = 0
    n = len(lines)

    while i < n:
        if i == title_line_idx:
            i += 1
            continue

        line    = lines[i]
        stripped = line.strip()

        # ── [CAPTION:] tag ────────────────────────────────────────────
        if stripped.startswith("[CAPTION:") and stripped.endswith("]"):
            pending_desc = stripped[9:-1].strip()
            i += 1
            continue

        # ── [DIAGRAM] block ───────────────────────────────────────────
        if stripped == "[DIAGRAM]":
            diagram_num += 1
            bk_id       += 1
            desc    = pending_desc or "Diagrama del sistema"
            anchor  = f"diag_{diagram_num}"
            caption_map[i] = ("Diagrama", desc, diagram_num, anchor, bk_id)
            pending_desc = None
            i += 1
            while i < n and lines[i].strip() != "[/DIAGRAM]":
                i += 1
            i += 1          # skip [/DIAGRAM]
            continue

        # ── Markdown table block ──────────────────────────────────────
        if "|" in line and stripped.startswith("|"):
            first_i = i
            while i < n and "|" in lines[i] and lines[i].strip().startswith("|"):
                i += 1
            table_num += 1
            bk_id     += 1
            desc   = pending_desc or "Datos tabulares"
            anchor = f"tbl_{table_num}"
            caption_map[first_i] = ("Tabla", desc, table_num, anchor, bk_id)
            table_toc_entries.append((table_num, desc, anchor))
            pending_desc = None
            continue

        # ── Fenced code block ─────────────────────────────────────────
        if line.startswith("```"):
            first_i = i
            i += 1
            while i < n and not lines[i].startswith("```"):
                i += 1
            i += 1          # skip closing ```
            code_num += 1
            bk_id    += 1
            desc   = pending_desc or "Fragmento de código"
            anchor = f"code_{code_num}"
            caption_map[first_i] = ("Código", desc, code_num, anchor, bk_id)
            pending_desc = None
            continue

        # ── Any other non-blank content clears pending ────────────────
        if stripped and not stripped.startswith("#"):
            pending_desc = None

        i += 1

    return caption_map, table_toc_entries, bk_id


# ── Inline markup parser ──────────────────────────────────────────────

_INLINE_RE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|`([^`]+)`|([^*`]+)", re.DOTALL)


def _add_inline(para, text: str, secondary_rgb: RGBColor) -> None:
    """Add *text* to *para*, honouring **bold**, *italic*, and `code`."""
    for m in _INLINE_RE.finditer(text):
        bold_txt, italic_txt, code_txt, plain_txt = m.groups()
        if bold_txt is not None:
            run = para.add_run(bold_txt)
            run.bold = True
        elif italic_txt is not None:
            run = para.add_run(italic_txt)
            run.italic = True
        elif code_txt is not None:
            run = para.add_run(code_txt)
            run.font.name      = "Courier New"
            run.font.size      = Pt(9)
            run.font.color.rgb = secondary_rgb
        elif plain_txt is not None:
            para.add_run(plain_txt)


# ── Document styling helpers ──────────────────────────────────────────

def _heading(doc: Document, text: str, level: int, first_h1: list,
             primary_rgb: RGBColor, secondary_rgb: RGBColor,
             anchor_id: str | None = None, bk_id: int | None = None) -> None:
    """Add a heading paragraph. If anchor_id/bk_id are given, attaches a bookmark."""
    if level == 1 and not first_h1:
        para = doc.add_paragraph(style="Title")
        para.clear()
        run = para.add_run(text)
        run.font.color.rgb = primary_rgb
        run.font.size = Pt(26)
        first_h1.append(True)
        if anchor_id is not None and bk_id is not None:
            _add_bookmark(para, anchor_id, bk_id)
        return

    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
    para  = doc.add_paragraph(style=style_map.get(level, "Heading 4"))
    para.clear()
    run   = para.add_run(text)
    run.font.color.rgb = primary_rgb if level <= 2 else secondary_rgb
    sizes = {1: Pt(20), 2: Pt(16), 3: Pt(13), 4: Pt(11)}
    run.font.size = sizes.get(level, Pt(11))

    if anchor_id is not None and bk_id is not None:
        _add_bookmark(para, anchor_id, bk_id)


def _code_block(doc: Document, lines: list[str], secondary_rgb: RGBColor) -> None:
    for line in lines:
        para = doc.add_paragraph(style="Normal")
        _set_para_shading(para, _COLOR_CODE_BG)
        para.paragraph_format.left_indent  = Inches(0.3)
        para.paragraph_format.right_indent = Inches(0.3)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        run = para.add_run(line)
        run.font.name      = "Courier New"
        run.font.size      = Pt(9)
        run.font.color.rgb = secondary_rgb


def _table_block(doc: Document, rows: list[str], primary_rgb: RGBColor,
                 primary_hex: str, secondary_rgb: RGBColor) -> None:
    def _split_row(line: str) -> list[str]:
        return [c.strip() for c in line.strip().strip("|").split("|")]

    data_rows = [r for r in rows if not re.match(r"^[\s|:\-]+$", r)]
    if not data_rows:
        return

    cols = len(_split_row(data_rows[0]))
    tbl  = doc.add_table(rows=len(data_rows), cols=cols)
    tbl.style = "Table Grid"
    _center_table(tbl)

    for r_idx, raw in enumerate(data_rows):
        cells = _split_row(raw)
        for c_idx, cell_text in enumerate(cells[:cols]):
            cell = tbl.cell(r_idx, c_idx)
            cell.text = ""
            para = cell.paragraphs[0]
            if r_idx == 0:
                _set_cell_bg(cell, primary_hex)
                run = para.add_run(cell_text)
                run.bold = True
                run.font.color.rgb = _COLOR_WHITE
            else:
                _add_inline(para, cell_text, secondary_rgb)

    doc.add_paragraph()


# ── Diagram block ────────────────────────────────────────────────────

def _inject_svg_into_inline(inline_shape, svg_bytes: bytes, doc: Document) -> None:
    """
    Extend an existing inline picture with an SVG alternative.

    Word 2016+ will display the SVG (vector, crisp at any zoom) and use the
    already-embedded PNG only as a fallback for older viewers.
    Users can right-click the diagram → "Convert to Shapes" to get fully
    editable native Word drawing objects.

    Parameters
    ----------
    inline_shape : docx.shape.InlineShape
        The result of ``run.add_picture()``.
    svg_bytes    : bytes
        Raw SVG data to embed.
    doc          : Document
        The parent document (needed to register the new relationship).
    """
    from docx.opc.part import Part as OpcPart
    from docx.opc.packuri import PackURI

    idx = next(_svg_counter)
    svg_part = OpcPart(
        PackURI(f"/word/media/diagram_svg_{idx}.svg"),
        "image/svg+xml",
        svg_bytes,
        doc.part.package,
    )
    rId_svg = doc.part.relate_to(svg_part, _REL_IMAGE)

    # Navigate wp:inline → a:graphic → a:graphicData → pic:pic → pic:blipFill → a:blip
    inline_el = inline_shape._inline
    graphic   = inline_el.find(qn("a:graphic"))
    if graphic is None:
        return
    graphic_data = graphic.find(qn("a:graphicData"))
    if graphic_data is None:
        return
    pic_el = graphic_data.find(qn("pic:pic"))
    if pic_el is None:
        return
    blip_fill = pic_el.find(qn("pic:blipFill"))
    if blip_fill is None:
        return
    blip = blip_fill.find(qn("a:blip"))
    if blip is None:
        return

    # Build: <a:extLst><a:ext uri="…"><asvg:svgBlip r:embed="rId"/></a:ext></a:extLst>
    from lxml import etree
    ext_xml = (
        f'<a:ext xmlns:a="{_NS_A}" '
        f'uri="{{96DAC541-7B7A-43D3-8B79-37D633B846F1}}">'
        f'<asvg:svgBlip xmlns:asvg="{_NS_ASVG}" '
        f'xmlns:r="{_NS_R}" '
        f'r:embed="{rId_svg}"/>'
        f'</a:ext>'
    )
    ext_el = etree.fromstring(ext_xml)

    ext_lst = blip.find(qn("a:extLst"))
    if ext_lst is None:
        ext_lst = OxmlElement("a:extLst")
        blip.append(ext_lst)
    ext_lst.append(ext_el)


def _diagram_block(doc: Document, mermaid_code: str,
                   primary_hex: str, secondary_rgb: RGBColor) -> None:
    """
    Insert a diagram into the document.

    Strategy
    --------
    1. Try to build a native DrawingML diagram (editable shapes + connectors).
       Supported for ``flowchart TD/LR`` and ``graph TD/LR`` only.
    2. If the diagram type is not supported by the native builder, or if the
       native build fails, fall back to fetching a PNG+SVG from mermaid.ink
       (three-attempt fallback chain inside diagram_renderer).
    3. If all rendering fails, insert an error note and the raw Mermaid code
       as a code block so the content is never silently lost.
    """
    from . import diagram_builder, diagram_renderer

    # ── 1. Native DrawingML (editable) ──────────────────────────────
    try:
        built = diagram_builder.build_native(doc, mermaid_code, primary_hex)
        if built:
            return
    except Exception:
        pass  # fall through to image rendering

    # ── 2. Image rendering (PNG) ─────────────────────────────────────
    png_bytes: bytes | None = None
    error_msg: str | None   = None

    try:
        png_bytes = diagram_renderer.render(mermaid_code, primary_hex)
    except Exception as exc:
        error_msg = str(exc)

    if png_bytes:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after  = Pt(8)
        para.add_run().add_picture(io.BytesIO(png_bytes), width=Inches(5.5))
        return

    # ── 3. Fallback: error note + raw Mermaid code ───────────────────
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after  = Pt(2)
    note_run = note.add_run(
        "[Diagrama no renderizado"
        + (f": {error_msg}" if error_msg else "")
        + " — código Mermaid:]"
    )
    note_run.font.italic    = True
    note_run.font.size      = Pt(9)
    note_run.font.color.rgb = _COLOR_GRAY
    _code_block(doc, mermaid_code.splitlines(), secondary_rgb)


# ── Cover page ────────────────────────────────────────────────────────

def _add_cover_page(doc: Document, title: str,
                    primary_rgb: RGBColor, secondary_rgb: RGBColor) -> None:
    month_year  = _month_year_es()
    body_sectPr = doc.element.body.sectPr
    orig_pgSz   = body_sectPr.find(qn("w:pgSz")) if body_sectPr is not None else None

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after  = Pt(20)
    run = title_para.add_run(title)
    run.font.size      = Pt(36)
    run.font.bold      = True
    run.font.color.rgb = primary_rgb

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_before = Pt(0)
    sub_para.paragraph_format.space_after  = Pt(30)
    sub_run = sub_para.add_run("Draft de documentación realizada por GenDoc")
    sub_run.font.size      = Pt(14)
    sub_run.font.italic    = True
    sub_run.font.color.rgb = secondary_rgb

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_before = Pt(0)
    date_para.paragraph_format.space_after  = Pt(4)
    date_run = date_para.add_run(month_year)
    date_run.font.size      = Pt(10)
    date_run.font.color.rgb = _COLOR_GRAY

    rights_para = doc.add_paragraph()
    rights_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rights_para.paragraph_format.space_before = Pt(0)
    rights_para.paragraph_format.space_after  = Pt(0)
    rights_run = rights_para.add_run("Todos los derechos reservados")
    rights_run.font.size      = Pt(10)
    rights_run.font.color.rgb = _COLOR_GRAY

    # Section break → next page (cover section with vAlign=center)
    sep_para = doc.add_paragraph()
    sep_para.paragraph_format.space_before = Pt(0)
    sep_para.paragraph_format.space_after  = Pt(0)

    pPr    = sep_para._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")

    type_el = OxmlElement("w:type")
    type_el.set(qn("w:val"), "nextPage")
    sectPr.append(type_el)

    if orig_pgSz is not None:
        sectPr.append(copy.deepcopy(orig_pgSz))
    else:
        pgSz = OxmlElement("w:pgSz")
        pgSz.set(qn("w:w"), "12240")
        pgSz.set(qn("w:h"), "15840")
        sectPr.append(pgSz)

    pgMar = OxmlElement("w:pgMar")
    for attr in ("w:top", "w:right", "w:bottom", "w:left", "w:header", "w:footer"):
        pgMar.set(qn(attr), "720")
    pgMar.set(qn("w:gutter"), "0")
    sectPr.append(pgMar)

    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    sectPr.append(vAlign)

    pPr.append(sectPr)


# ── TOC / Agenda page ────────────────────────────────────────────────

def _add_toc_page(
    doc: Document,
    headings: list[tuple[int, str, str]],
    toc_title: str,
    primary_rgb:       RGBColor,
    primary_hex:       str,
    secondary_rgb:     RGBColor,
    secondary_hex:     str,
    table_toc_entries: list[tuple[int, str, str]] | None = None,
) -> None:
    """
    Add a TOC/Agenda page after the cover page.

    *headings* → (level, text, anchor_id) for H1–H3.
    *table_toc_entries* → (number, description, anchor_id) for the Índice de Tablas.
    Each entry is a clickable internal hyperlink.
    A page break at the end forces body content onto the next page.
    """
    # ── Section title ────────────────────────────────────────────────
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc_heading.paragraph_format.space_before = Pt(0)
    toc_heading.paragraph_format.space_after  = Pt(16)
    run = toc_heading.add_run(toc_title)
    run.font.size      = Pt(20)
    run.font.bold      = True
    run.font.color.rgb = primary_rgb

    # ── Heading entries ───────────────────────────────────────────────
    _LEVEL_CFG = {
        1: (Inches(0.00), 22, True,  primary_hex),
        2: (Inches(0.30), 20, False, primary_hex),
        3: (Inches(0.55), 18, False, secondary_hex),
    }

    for level, text, anchor in headings:
        indent, half_pts, bold, hex_col = _LEVEL_CFG.get(
            level, (Inches(0.55), 18, False, secondary_hex)
        )
        para = doc.add_paragraph()
        para.paragraph_format.left_indent  = indent
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)
        _add_toc_hyperlink(para, text, anchor, hex_col, half_pts, bold)

    # ── Índice de Tablas ──────────────────────────────────────────────
    if table_toc_entries:
        # Spacer
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(14)
        spacer.paragraph_format.space_after  = Pt(8)
        run_title = spacer.add_run("Índice de Tablas")
        run_title.font.size      = Pt(13)
        run_title.font.bold      = True
        run_title.font.color.rgb = primary_rgb

        for num, description, anchor in table_toc_entries:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent  = Inches(0.20)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            _add_toc_hyperlink(
                para,
                f"Tabla {num},  {description}",
                anchor,
                primary_hex,
                20,
                False,
            )

    # ── Page break → body starts on next page ────────────────────────
    _add_page_break(doc)


# ── Header & footer ───────────────────────────────────────────────────

def _setup_header_footer(doc: Document, project_title: str,
                         primary_rgb: RGBColor, secondary_rgb: RGBColor) -> None:
    month_year = _month_year_es()
    section    = doc.sections[-1]

    # ── Header (borderless 2-column table) ───────────────────────────
    header  = section.header
    hdr_tbl = header.add_table(rows=1, cols=2, width=Cm(19.05))
    _remove_table_borders(hdr_tbl)

    _set_cell_width(hdr_tbl.cell(0, 0), 1872)   # ~1.3 in
    _set_cell_width(hdr_tbl.cell(0, 1), 8928)   # rest

    cell_img = hdr_tbl.cell(0, 0)
    _set_cell_valign(cell_img, "center")
    p_img = cell_img.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_img = p_img.add_run()
    run_img.add_picture(io.BytesIO(_make_placeholder_png(120, 50)), width=Inches(1.05))

    cell_txt = hdr_tbl.cell(0, 1)
    _set_cell_valign(cell_txt, "center")

    p_author = cell_txt.paragraphs[0]
    p_author.paragraph_format.space_before = Pt(0)
    p_author.paragraph_format.space_after  = Pt(1)
    r_author = p_author.add_run("Autor(a):")
    r_author.font.size      = Pt(9)
    r_author.font.bold      = True
    r_author.font.color.rgb = primary_rgb

    p_draft = cell_txt.add_paragraph()
    p_draft.paragraph_format.space_before = Pt(0)
    p_draft.paragraph_format.space_after  = Pt(0)
    r_draft = p_draft.add_run("Draft generado por GenDoc")
    r_draft.font.size      = Pt(8)
    r_draft.font.italic    = True
    r_draft.font.color.rgb = primary_rgb

    hdr_el  = header._element
    first_p = hdr_el.find(qn("w:p"))
    if first_p is not None:
        hdr_el.remove(first_p)
    header.add_paragraph()

    # ── Footer ───────────────────────────────────────────────────────
    footer = section.footer
    fp     = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    r_title = fp.add_run(project_title)
    r_title.font.bold      = True
    r_title.font.size      = Pt(9)
    r_title.font.color.rgb = primary_rgb

    r_sep = fp.add_run("   ·   ")
    r_sep.font.size      = Pt(9)
    r_sep.font.color.rgb = _COLOR_GRAY

    r_rights = fp.add_run(f"Todos los derechos reservados, {month_year}")
    r_rights.font.size      = Pt(9)
    r_rights.font.color.rgb = _COLOR_GRAY


# ── Main converter ────────────────────────────────────────────────────

def convert(
    markdown: str,
    output_path: str | Path,
    *,
    doc_type:        str = "technical",
    primary_color:   str = _DEFAULT_PRIMARY,
    secondary_color: str = _DEFAULT_SECONDARY,
) -> Path:
    """
    Convert *markdown* to a styled .docx and write it to *output_path*.

    Parameters
    ----------
    doc_type        : 'technical' | 'user_manual' | 'executive'
    primary_color   : hex string — title, H1/H2, table headers, header/footer text
    secondary_color : hex string — H3+, code elements

    Raises
    ------
    RuntimeError — if the parent directory does not exist or is not writable.
    """
    out = Path(output_path)
    if not out.parent.exists():
        raise RuntimeError(
            f"El directorio de salida no existe: {out.parent}. "
            "Crea la carpeta o cambia OUTPUT_DIR en .env."
        )
    if not out.parent.is_dir():
        raise RuntimeError(f"La ruta de salida no es una carpeta: {out.parent}.")

    primary_rgb   = _hex_to_rgb(primary_color)
    primary_hex   = _hex_to_str(primary_color)
    secondary_rgb = _hex_to_rgb(secondary_color)
    secondary_hex = _hex_to_str(secondary_color)

    doc = Document()

    # ── Default body font ────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = markdown.splitlines()
    n     = len(lines)

    # ── Find first H1 → cover title ──────────────────────────────────
    title          = None
    title_line_idx = -1
    for idx, line in enumerate(lines):
        m = re.match(r"^#\s+(.*)", line)
        if m:
            title          = m.group(1).strip()
            title_line_idx = idx
            break
    if not title:
        title = out.stem.replace("_", " ").title()

    # ── Pre-scan 1: headings (for TOC links + body bookmarks) ────────
    heading_map: dict[int, tuple[str, int]] = {}
    toc_entries: list[tuple[int, str, str]] = []
    seen_anchors: dict[str, int] = {}
    bk_counter = 1

    for idx, line in enumerate(lines):
        if idx == title_line_idx:
            continue
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if not m:
            continue
        level = len(m.group(1))
        text  = m.group(2).strip()
        slug  = re.sub(r"[^\w]", "_", text)[:40].strip("_")
        if not slug or slug[0].isdigit():
            slug = "h_" + slug
        count = seen_anchors.get(slug, 0)
        seen_anchors[slug] = count + 1
        anchor = slug if count == 0 else f"{slug}_{count}"
        heading_map[idx] = (anchor, bk_counter)
        if level <= 3:
            toc_entries.append((level, text, anchor))
        bk_counter += 1

    # ── Pre-scan 2: tables / code / diagrams (for captions + table TOC) ─
    caption_map, table_toc_entries, bk_counter = _prescan_elements(
        lines, title_line_idx, bk_counter
    )

    # ── Cover page (section 1, vAlign=center) ────────────────────────
    _add_cover_page(doc, title, primary_rgb, secondary_rgb)

    # ── Header + footer on main section ──────────────────────────────
    _setup_header_footer(doc, title, primary_rgb, secondary_rgb)

    # ── Apply 1.27 cm margins to all sections ────────────────────────
    for sec in doc.sections:
        sec.top_margin      = Cm(1.27)
        sec.bottom_margin   = Cm(1.27)
        sec.left_margin     = Cm(1.27)
        sec.right_margin    = Cm(1.27)
        sec.header_distance = Cm(1.27)
        sec.footer_distance = Cm(1.27)

    # ── TOC / Agenda page ────────────────────────────────────────────
    toc_title = "Agenda" if doc_type == "executive" else "Índice"
    if toc_entries or table_toc_entries:
        _add_toc_page(doc, toc_entries, toc_title,
                      primary_rgb, primary_hex, secondary_rgb, secondary_hex,
                      table_toc_entries)

    # ── Process markdown body ─────────────────────────────────────────
    # first_h1 pre-filled: cover consumed the title H1, remaining H1s
    # render as Heading 1 (not the Title style).
    first_h1 = [True]
    i        = 0

    while i < n:
        line = lines[i]

        if i == title_line_idx:
            i += 1
            continue

        # ── [CAPTION:] tag — skip (already incorporated in caption_map) ──
        if lines[i].strip().startswith("[CAPTION:") and lines[i].strip().endswith("]"):
            i += 1
            continue

        # ── Diagram block ────────────────────────────────────────────
        if line.strip() == "[DIAGRAM]":
            diag_start    = i
            diagram_lines: list[str] = []
            i += 1
            while i < n and lines[i].strip() != "[/DIAGRAM]":
                diagram_lines.append(lines[i])
                i += 1
            _diagram_block(doc, "\n".join(diagram_lines),
                           primary_hex, secondary_rgb)
            cap = caption_map.get(diag_start)
            if cap:
                _add_caption(doc, cap[0], cap[2], cap[1], cap[3], cap[4])
            i += 1  # skip [/DIAGRAM]
            continue

        # ── Fenced code block ────────────────────────────────────────
        if line.startswith("```"):
            code_start  = i
            code_lines: list[str] = []
            i += 1
            while i < n and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            _code_block(doc, code_lines, secondary_rgb)
            cap = caption_map.get(code_start)
            if cap:
                _add_caption(doc, cap[0], cap[2], cap[1], cap[3], cap[4])
            i += 1
            continue

        # ── Markdown table ───────────────────────────────────────────
        if "|" in line and line.strip().startswith("|"):
            tbl_start  = i
            table_rows: list[str] = []
            while i < n and "|" in lines[i] and lines[i].strip().startswith("|"):
                table_rows.append(lines[i])
                i += 1
            _table_block(doc, table_rows, primary_rgb, primary_hex, secondary_rgb)
            cap = caption_map.get(tbl_start)
            if cap:
                _add_caption(doc, cap[0], cap[2], cap[1], cap[3], cap[4])
            continue

        # ── Headings ─────────────────────────────────────────────────
        heading_m = re.match(r"^(#{1,4})\s+(.*)", line)
        if heading_m:
            level     = len(heading_m.group(1))
            text      = heading_m.group(2).strip()
            anchor_id, bk_id = heading_map.get(i, (None, None))
            _heading(doc, text, level, first_h1, primary_rgb, secondary_rgb,
                     anchor_id, bk_id)
            i += 1
            continue

        # ── Unordered list ───────────────────────────────────────────
        if re.match(r"^\s*[-*]\s+", line):
            para = doc.add_paragraph(style="List Bullet")
            _add_inline(para, re.sub(r"^\s*[-*]\s+", "", line), secondary_rgb)
            i += 1
            continue

        # ── Ordered list ─────────────────────────────────────────────
        if re.match(r"^\s*\d+\.\s+", line):
            para = doc.add_paragraph(style="List Number")
            _add_inline(para, re.sub(r"^\s*\d+\.\s+", "", line), secondary_rgb)
            i += 1
            continue

        # ── Horizontal rule ──────────────────────────────────────────
        if re.match(r"^---+$", line.strip()):
            doc.add_paragraph()
            i += 1
            continue

        # ── Blank line ───────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Normal paragraph ─────────────────────────────────────────
        para = doc.add_paragraph(style="Normal")
        _add_inline(para, line, secondary_rgb)
        i += 1

    # ── Write file ───────────────────────────────────────────────────
    try:
        doc.save(str(out))
    except PermissionError:
        raise RuntimeError(
            f"Sin permisos para escribir en: {out}. "
            "Cierra el archivo si está abierto en Word e intenta de nuevo."
        )
    except OSError as exc:
        raise RuntimeError(f"No se pudo guardar el documento Word: {exc}")

    return out
